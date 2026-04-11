"""
Улучшенный File Processor для ReplyMan

Поддерживаемые форматы:
- txt, json, html (извлечение текста + оптимизация)
- pdf (извлечение текста через PyMuPDF)
- docx (извлечение текста через python-docx)
- doc (конвертация через libreoffice → docx → текст)
"""

from bs4 import BeautifulSoup
import json
from typing import Dict, Optional, List
import re
from datetime import datetime
import io
import os
import tempfile
import subprocess
import logging

logger = logging.getLogger(__name__)


class TelegramOptimizer:
    """Встроенный оптимизатор для Telegram JSON экспортов"""

    FIELDS_TO_REMOVE = {
        'id', 'message_id', 'reply_to_message_id',
        'photo', 'photo_file_size', 'width', 'height',
        'file', 'file_name', 'file_size', 'thumbnail', 'thumbnail_file_size',
        'media_type', 'mime_type', 'duration_seconds',
        'date_unixtime', 'from_id', 'text_entities',
        'type', 'edited', 'edited_unixtime',
        'forwarded_from', 'via_bot', 'actor', 'actor_id',
    }

    ROOT_FIELDS_TO_REMOVE = {
        'about', 'personal_information', 'profile_pictures',
        'frequent_contacts', 'user_info'
    }


class FileProcessor:
    """Service for processing uploaded files (txt, json, html, pdf, docx, doc)"""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.last_stats = {}

    async def process_file(self, content: bytes, content_type: str, filename: str) -> Dict:
        """Process uploaded file and extract text content"""
        try:
            name_lower = filename.lower()

            # PDF
            if name_lower.endswith('.pdf') or 'pdf' in content_type:
                return self._process_pdf(content)

            # DOCX
            if name_lower.endswith('.docx') or 'wordprocessingml' in content_type:
                return self._process_docx(content)

            # DOC (старый формат)
            if name_lower.endswith('.doc') and not name_lower.endswith('.docx'):
                return await self._process_doc(content)

            # HTML
            if name_lower.endswith(('.html', '.htm')) or 'html' in content_type:
                text_content = content.decode('utf-8', errors='ignore')
                return self._process_html(text_content)

            # JSON
            if name_lower.endswith('.json') or 'json' in content_type:
                text_content = content.decode('utf-8', errors='ignore')
                return self._process_json(text_content)

            # TXT / всё остальное — как обычный текст
            text_content = content.decode('utf-8', errors='ignore')
            return self._process_txt(text_content)

        except Exception as e:
            logger.error(f"process_file error for {filename}: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "content": ""
            }

    # ========================================
    # PDF
    # ========================================

    def _process_pdf(self, content: bytes) -> Dict:
        """Извлечь текст из PDF через PyMuPDF (fitz)"""
        original_size = len(content)
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text.strip())

            doc.close()

            full_text = "\n\n".join(text_parts)
            cleaned = self._clean_text(full_text)

            return {
                "success": True,
                "content": cleaned,
                "conversations": [],
                "format": "pdf",
                "stats": {
                    "original_size": original_size,
                    "optimized_size": len(cleaned),
                    "pages": len(text_parts),
                    "compression_ratio": round((1 - len(cleaned) / max(original_size, 1)) * 100, 1)
                }
            }
        except ImportError:
            logger.error("PyMuPDF (fitz) не установлен. Установите: pip install PyMuPDF")
            return {
                "success": False,
                "error": "Библиотека PyMuPDF не установлена. pip install PyMuPDF",
                "content": ""
            }
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {
                "success": False,
                "error": f"Ошибка чтения PDF: {e}",
                "content": ""
            }

    # ========================================
    # DOCX
    # ========================================

    def _process_docx(self, content: bytes) -> Dict:
        """Извлечь текст из DOCX через python-docx"""
        original_size = len(content)
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            cleaned = self._clean_text(full_text)

            return {
                "success": True,
                "content": cleaned,
                "conversations": [],
                "format": "docx",
                "stats": {
                    "original_size": original_size,
                    "optimized_size": len(cleaned),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "compression_ratio": round((1 - len(cleaned) / max(original_size, 1)) * 100, 1)
                }
            }
        except ImportError:
            logger.error("python-docx не установлен. Установите: pip install python-docx")
            return {
                "success": False,
                "error": "Библиотека python-docx не установлена. pip install python-docx",
                "content": ""
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {
                "success": False,
                "error": f"Ошибка чтения DOCX: {e}",
                "content": ""
            }

    # ========================================
    # DOC (старый формат Word)
    # ========================================

    async def _process_doc(self, content: bytes) -> Dict:
        """Извлечь текст из DOC: пробуем DOCX → antiword → LibreOffice → textract"""
        original_size = len(content)

        # Шаг 1: Пробуем прочитать как DOCX (некоторые .doc на самом деле DOCX)
        try:
            result = self._process_docx(content)
            if result.get("success"):
                logger.info("Файл .doc оказался DOCX, прочитан успешно")
                result["format"] = "doc"
                return result
        except Exception:
            pass  # Не DOCX, продолжаем

        # Шаг 2: Пробуем antiword (быстрый, но не всегда установлен)
        antiword_result = await self._try_antiword(content, original_size)
        if antiword_result:
            return antiword_result

        # Шаг 3: Пробуем LibreOffice (надёжный, но медленный)
        libreoffice_result = await self._try_libreoffice_doc(content, original_size)
        if libreoffice_result:
            return libreoffice_result

        # Шаг 4: Пробуем textract (Python-библиотека)
        textract_result = self._try_textract(content, original_size)
        if textract_result:
            return textract_result

        return {
            "success": False,
            "error": "Не удалось прочитать DOC файл. Установите antiword (apt install antiword) или LibreOffice (apt install libreoffice).",
            "content": ""
        }

    async def _try_antiword(self, content: bytes, original_size: int) -> Optional[Dict]:
        """Попытка извлечь текст через antiword"""
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            try:
                antiword_bin = "/usr/bin/antiword"
                if not os.path.exists(antiword_bin):
                    antiword_bin = "antiword"  # fallback

                result = subprocess.run(
                    [antiword_bin, tmp_path],
                    capture_output=True,
                    timeout=30
                )

                if result.returncode != 0:
                    logger.warning(f"antiword failed: {result.stderr.decode(errors='ignore').strip()}")
                    return None

                text = result.stdout.decode('utf-8', errors='ignore')
                cleaned = self._clean_text(text)

                if not cleaned.strip():
                    return None

                logger.info(f"DOC extracted via antiword: {len(cleaned)} chars")
                return {
                    "success": True,
                    "content": cleaned,
                    "conversations": [],
                    "format": "doc",
                    "stats": {
                        "original_size": original_size,
                        "optimized_size": len(cleaned),
                        "compression_ratio": round((1 - len(cleaned) / max(original_size, 1)) * 100, 1)
                    }
                }
            finally:
                os.unlink(tmp_path)

        except FileNotFoundError:
            logger.info("antiword не найден в системе, пробуем другие методы")
            return None
        except subprocess.TimeoutExpired:
            logger.warning("antiword таймаут")
            return None
        except Exception as e:
            logger.warning(f"antiword error: {e}")
            return None

    async def _try_libreoffice_doc(self, content: bytes, original_size: int) -> Optional[Dict]:
        """Попытка извлечь текст через LibreOffice (конвертация в txt)"""
        tmp_path = None
        tmp_dir = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            tmp_dir = tempfile.mkdtemp(prefix="replyman_doc_")

            # Конвертируем в txt
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmp_dir, tmp_path],
                capture_output=True,
                timeout=60
            )

            if result.returncode != 0:
                logger.warning(f"LibreOffice conversion failed: {result.stderr.decode(errors='ignore').strip()}")
                return None

            # Ищем созданный .txt файл
            txt_files = [f for f in os.listdir(tmp_dir) if f.endswith('.txt')]
            if not txt_files:
                logger.warning("LibreOffice не создал txt файл")
                return None

            txt_path = os.path.join(tmp_dir, txt_files[0])
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            cleaned = self._clean_text(text)

            if not cleaned.strip():
                return None

            logger.info(f"DOC extracted via LibreOffice: {len(cleaned)} chars")
            return {
                "success": True,
                "content": cleaned,
                "conversations": [],
                "format": "doc",
                "stats": {
                    "original_size": original_size,
                    "optimized_size": len(cleaned),
                    "compression_ratio": round((1 - len(cleaned) / max(original_size, 1)) * 100, 1)
                }
            }

        except FileNotFoundError:
            logger.info("LibreOffice не найден в системе")
            return None
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice таймаут (60 секунд)")
            return None
        except Exception as e:
            logger.warning(f"LibreOffice DOC error: {e}")
            return None
        finally:
            # Удаляем временные файлы
            try:
                if tmp_path and os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                if tmp_dir and os.path.exists(tmp_dir):
                    import shutil
                    shutil.rmtree(tmp_dir)
            except Exception:
                pass

    def _try_textract(self, content: bytes, original_size: int) -> Optional[Dict]:
        """Попытка извлечь текст через textract (Python-библиотека)"""
        try:
            import textract

            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            try:
                text = textract.process(tmp_path).decode('utf-8', errors='ignore')
                cleaned = self._clean_text(text)

                if not cleaned.strip():
                    return None

                logger.info(f"DOC extracted via textract: {len(cleaned)} chars")
                return {
                    "success": True,
                    "content": cleaned,
                    "conversations": [],
                    "format": "doc",
                    "stats": {
                        "original_size": original_size,
                        "optimized_size": len(cleaned),
                        "compression_ratio": round((1 - len(cleaned) / max(original_size, 1)) * 100, 1)
                    }
                }
            finally:
                os.unlink(tmp_path)

        except ImportError:
            logger.info("textract не установлен")
            return None
        except Exception as e:
            logger.warning(f"textract error: {e}")
            return None

    # ========================================
    # TXT
    # ========================================

    def _process_txt(self, content: str) -> Dict:
        """Process plain text file"""
        cleaned = self._clean_text(content)
        conversations = self._parse_text_conversation(cleaned)

        return {
            "success": True,
            "content": cleaned,
            "conversations": conversations,
            "format": "txt",
            "stats": {
                "original_size": len(content),
                "optimized_size": len(cleaned),
                "compression_ratio": round((1 - len(cleaned) / max(len(content), 1)) * 100, 1)
            }
        }

    # ========================================
    # JSON
    # ========================================

    def _process_json(self, content: str) -> Dict:
        """Process JSON file with Telegram optimization"""
        original_size = len(content)

        try:
            data = json.loads(content)

            if isinstance(data, list):
                result = self._process_json_messages(data)
            elif isinstance(data, dict):
                result = self._process_json_dict(data)
            else:
                result = {
                    "success": True,
                    "content": json.dumps(data, ensure_ascii=False),
                    "conversations": [],
                    "format": "json"
                }

            optimized_size = len(result.get("content", ""))
            result["stats"] = {
                "original_size": original_size,
                "optimized_size": optimized_size,
                "compression_ratio": round((1 - optimized_size / max(original_size, 1)) * 100, 1)
            }

            return result

        except json.JSONDecodeError:
            return self._process_txt(content)

    def _process_json_messages(self, messages: List) -> Dict:
        """Process JSON array of messages"""
        conversations = []
        text_parts = []

        for msg in messages:
            if isinstance(msg, dict):
                sender = msg.get('sender', msg.get('from', msg.get('author', msg.get('name', 'Unknown'))))
                text = self._extract_text(msg.get('text', msg.get('message', msg.get('content', msg.get('body', '')))))
                timestamp = msg.get('timestamp', msg.get('date', msg.get('time', '')))

                if not text or not text.strip():
                    continue

                if timestamp:
                    timestamp = self._shorten_date(str(timestamp))

                line = f"[{sender}]: {text}"
                if timestamp:
                    line = f"{timestamp} {line}"

                text_parts.append(line)
                conversations.append({
                    "sender": sender,
                    "text": text,
                    "timestamp": timestamp
                })

        return {
            "success": True,
            "content": "\n".join(text_parts),
            "conversations": conversations,
            "format": "json_messages"
        }

    def _process_json_dict(self, data: Dict) -> Dict:
        """Process JSON dictionary - оптимизированная версия для Telegram"""
        conversations = []
        text_parts = []
        stats = {
            "messages_total": 0,
            "messages_with_text": 0,
            "messages_empty": 0,
            "chats_count": 0
        }

        if 'chats' in data:
            chats = data['chats'].get('list', [])
            stats["chats_count"] = len(chats)

            for chat in chats:
                chat_name = chat.get('name', 'Чат')
                messages = chat.get('messages', [])

                if messages:
                    text_parts.append(f"\n=== {chat_name} ===")

                for msg in messages:
                    stats["messages_total"] += 1

                    if msg.get('type') not in ['message', None]:
                        continue

                    sender = msg.get('from', 'Unknown')
                    text = self._extract_text(msg.get('text', ''))
                    date = msg.get('date', '')

                    if not text or not text.strip():
                        stats["messages_empty"] += 1
                        continue

                    stats["messages_with_text"] += 1

                    if date:
                        date = self._shorten_date(str(date))

                    if date:
                        line = f"{date} [{sender}]: {text}"
                    else:
                        line = f"[{sender}]: {text}"

                    text_parts.append(line)
                    conversations.append({
                        "sender": sender,
                        "text": text,
                        "timestamp": date,
                        "chat": chat_name
                    })

        elif 'messages' in data:
            messages = data['messages']
            if isinstance(messages, list):
                stats["messages_total"] = len(messages)

                for msg in messages:
                    sender = msg.get('sender', msg.get('from', 'Unknown'))
                    text = msg.get('text', msg.get('message', ''))

                    if text:
                        text = self._extract_text(text)
                        if text.strip():
                            stats["messages_with_text"] += 1
                            text_parts.append(f"[{sender}]: {text}")
                            conversations.append({
                                "sender": sender,
                                "text": str(text),
                                "timestamp": msg.get('timestamp', '')
                            })

        else:
            def extract_messages(obj):
                if isinstance(obj, dict):
                    if 'message' in obj or 'text' in obj:
                        sender = obj.get('sender', obj.get('from', obj.get('author', 'Unknown')))
                        text = obj.get('message', obj.get('text', ''))
                        if text:
                            text = self._extract_text(text)
                            if text.strip():
                                text_parts.append(f"[{sender}]: {text}")
                                conversations.append({
                                    "sender": sender,
                                    "text": str(text),
                                    "timestamp": obj.get('timestamp', obj.get('date', ''))
                                })
                    else:
                        for value in obj.values():
                            extract_messages(value)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_messages(item)

            extract_messages(data)

        self.last_stats = stats

        return {
            "success": True,
            "content": "\n".join(text_parts),
            "conversations": conversations,
            "format": "telegram_optimized",
            "messages_stats": stats
        }

    # ========================================
    # HTML
    # ========================================

    def _process_html(self, content: str) -> Dict:
        """Process HTML file"""
        soup = BeautifulSoup(content, 'lxml')
        conversations = []
        text_parts = []

        for script in soup(["script", "style"]):
            script.decompose()

        chat_containers = soup.find_all(['div'], class_=re.compile(r'(message|chat|msg)', re.I))

        if chat_containers:
            for container in chat_containers:
                text = container.get_text(strip=True, separator=' ')
                if text and len(text) > 5:
                    text_parts.append(text)
                    conversations.append({"text": text})
        else:
            for elem in soup.find_all(['p', 'li', 'div']):
                text = elem.get_text(strip=True, separator=' ')
                if text and len(text) > 10:
                    text_parts.append(text)

        plain_text = soup.get_text(separator='\n', strip=True)

        return {
            "success": True,
            "content": plain_text if plain_text else "\n".join(text_parts),
            "conversations": conversations,
            "format": "html",
            "stats": {
                "original_size": len(content),
                "optimized_size": len(plain_text),
                "compression_ratio": round((1 - len(plain_text) / max(len(content), 1)) * 100, 1)
            }
        }

    # ========================================
    # Вспомогательные методы
    # ========================================

    def _extract_text(self, text_field) -> str:
        """Извлечь текст из поля (строка или массив с форматированием)."""
        if isinstance(text_field, str):
            return text_field

        if isinstance(text_field, list):
            parts = []
            for item in text_field:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, dict):
                    text = item.get('text', '')
                    item_type = item.get('type', '')
                    if item_type == 'phone':
                        text = f"📞{text}"
                    elif item_type == 'email':
                        text = f"📧{text}"
                    elif item_type == 'link':
                        text = f"🔗{text}"
                    parts.append(text)
            return ''.join(parts)

        return str(text_field) if text_field else ''

    def _shorten_date(self, date_str: str) -> str:
        """Сократить формат даты: 2025-01-15T11:44:49 -> 15.01 11:44"""
        try:
            date_str = date_str.replace('Z', '')
            if 'T' in date_str:
                dt = datetime.fromisoformat(date_str.split('+')[0])
                return dt.strftime('%d.%m %H:%M')
            else:
                return date_str
        except:
            return date_str

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()

    def _parse_text_conversation(self, text: str) -> List[Dict]:
        """Parse conversation format from plain text"""
        conversations = []

        patterns = [
            r'^(\d{1,2}[./]\d{1,2}[./]\d{2,4}.*?)$',
            r'^\[?(\d{1,2}:\d{2})\]?.*?:',
            r'^([^:]+):\s*(.+)$',
            r'^<([^>]+)>\s*(.+)$',
        ]

        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    groups = match.groups()
                    if len(groups) >= 2:
                        conversations.append({
                            "sender": groups[0] if groups[0] else "Unknown",
                            "text": groups[1] if len(groups) > 1 else line,
                            "raw": line
                        })
                        break
            else:
                if conversations:
                    conversations[-1]["text"] += f" {line}"

        return conversations

    def estimate_tokens(self, text: str) -> int:
        """Оценить количество токенов (~4 символа на токен для русского)"""
        return len(text) // 4


# Singleton instance
file_processor = FileProcessor()